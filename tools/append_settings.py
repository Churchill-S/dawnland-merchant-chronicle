# -*- coding: utf-8 -*-
"""把苍澜洲/航运公司章节追加进两份设定 docx（按文件大小识别，避免中文文件名编码问题）"""
import glob
import os
import sys

import docx


def find(size):
    hits = [p for p in glob.glob(r"D:\Churchill\Dawnland\*.docx") if os.path.getsize(p) == size]
    return hits[0] if hits else None


def append_world(path):
    doc = docx.Document(path)
    doc.add_heading("补充章：苍澜洲（海外新大陆·游戏化扩展）", level=1)
    doc.add_paragraph(
        "曦光之地西南海外另有一块大陆，名苍澜洲。这里海雾常年不散、雨水丰沛，居民以渔、盐、造船与海贸为生，"
        "组成城邦联盟苍澜海盟，货币为苍澜贝币。"
    )
    doc.add_heading("城市（7座）", level=2)
    t = doc.add_table(rows=8, cols=3)
    try:
        t.style = "Table Grid"
    except Exception:
        pass
    rows = [
        ("城市", "类型", "特色"),
        ("沧浪港", "大港·首府", "稀有咸鱼、精良花香盐与木雕"),
        ("鲸歌湾", "渔港", "稀有咸鱼、精良橄榄油，海湾常有鲸群"),
        ("雾灯屿", "灯塔岛", "稀有药草、精良松脂糖，常年海雾"),
        ("礁石城", "渔镇", "稀有彩陶（珊瑚陶）、精良咸鱼，环礁密布"),
        ("风帆镇", "船坞镇", "稀有木雕、精良麦酒，造船匠之乡"),
        ("盐沫镇", "盐田", "稀有花香盐"),
        ("青霭城", "内陆都会", "稀有药草、精良书籍（海图），雾绕钟楼"),
    ]
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            t.rows[i].cells[j].text = v
    doc.add_heading("海上航路（远澜航线）", level=2)
    doc.add_paragraph(
        "曦光至苍澜主干航线：青贝港至沧浪港8天、盐花原至雾灯屿6天、金河三角洲至鲸歌湾7天、白石渡至风帆镇6天。"
        "苍澜洲内部航线：沧浪港至鲸歌湾3天、沧浪港至雾灯屿2天、雾灯屿至礁石城3天、风帆镇至鲸歌湾2天。"
        "内陆道路：沧青官道、青盐驿道、盐礁小径、帆盐商道。"
        "远洋航线使青贝港、盐花原、白石渡、金河三角洲等港口城市价值大幅提升。"
    )
    doc.save(path)
    print("已追加苍澜洲章节：", path)


def append_game(path):
    doc = docx.Document(path)
    doc.add_heading("十二、补充：航运公司（新增建筑）", level=1)
    doc.add_paragraph(
        "港口城市（临海）可建造航运公司，费用700G、工期15天，可升级至3级。"
        "收入按本城近30天海上商路的使用次数计算航运费（用船越多赚得越多）；"
        "每日需支付船只维护费；如果海路冷清、无人需要船只，则只有维护费支出而没有进账。"
    )
    doc.add_heading("地图交互调整", level=2)
    doc.add_paragraph("地图改为左键拖拽平移，点击城镇查看详情；滚轮缩放不变。")
    doc.save(path)
    print("已追加航运公司章节：", path)


def main():
    sys.stdout.reconfigure(encoding="utf-8")
    world = find(49927)
    game = find(21486)
    if world:
        append_world(world)
    else:
        print("未找到世界观设定.docx（当前大小不符）")
    if game:
        append_game(game)
    else:
        print("未找到游戏设定.docx（当前大小不符）")


if __name__ == "__main__":
    main()
